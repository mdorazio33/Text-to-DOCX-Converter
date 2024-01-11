import fitz  # PyMuPDF
import spacy
from docx import Document
import easyocr
import requests
from bs4 import BeautifulSoup
import numpy as np
import warnings
from bs4 import Tag
from PIL import Image, ImageEnhance
from io import BytesIO
import cv2

warnings.simplefilter(action='ignore', category=UserWarning)

def is_text_only_page(page):
    # Code to check if the page contains only text (no images)
    return not page.get_images(full=True)

def is_image_page(page):
    # Code to check if the page contains images
    return bool(page.get_images(full=True))

def pdf_to_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    has_images = False

    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()

        # Code to check if the page contains images
        if not has_images and is_image_page(page):
            has_images = True

    doc.close()
    return text, has_images

def extract_images_from_pdf(pdf_path):
    images = []
    doc = fitz.open(pdf_path)

    for page_num in range(doc.page_count):
        page = doc[page_num]
        images += page.get_images(full=True)

    doc.close()
    return images

def preprocess_image_opencv(image, resize_factor=10, desired_dpi=300):
    # Code to resize the image by a factor of 10
    resized_image = cv2.resize(image, None, fx=resize_factor, fy=resize_factor, interpolation=cv2.INTER_CUBIC)

    # Code to convert the image to grayscale
    gray_image = cv2.cvtColor(resized_image, cv2.COLOR_BGR2GRAY)

    # Code to apply a Gaussian blur to reduce noise
    blurred_image = cv2.GaussianBlur(gray_image, (5, 5), 0)

    # Code to apply adaptive thresholding to enhance contrast
    _, threshold_image = cv2.threshold(blurred_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Code to apply morphological transformations to further clean the image
    kernel = np.ones((3, 3), np.uint8)
    processed_image = cv2.morphologyEx(threshold_image, cv2.MORPH_CLOSE, kernel, iterations=2)

    # Code to adjust the intensity of black pixels in the processed image
    processed_image = cv2.subtract(255, processed_image)

    return processed_image

def extract_text_from_image(image):
    # Code used to load the preprocessed image
    preprocessed_image = preprocess_image_opencv(image)

    # Code used to enable easyocr for text extraction
    reader = easyocr.Reader(['en'])
    result = reader.readtext(preprocessed_image)

    # Code used to extract text from the result
    text = ' '.join([entry[1] for entry in result])
    return text

def extract_text_from_website(url):
    # Code used to fetch HTML content from the website
    response = requests.get(url)
    html_content = response.text

    # Code used to enable BeautifulSoup to parse HTML and extract text
    soup = BeautifulSoup(html_content, 'html.parser')

    def process_list_items(list_tag, level=0):
        # Recursively process list items and maintain indentation
        items = list_tag.find_all('li', recursive=False)
        tags = []
        for index, item in enumerate(items, start=1):
            text = f"{'  ' * level}{'• ' if list_tag.name == 'ul' else f'{index}. '}{item.get_text(strip=True)}"
            if item.find_all(['ul', 'ol']):
                tags.extend(process_list_items(item, level + 1))
            tag = Tag(name='text')
            tag.string = text
            tag.level = level
            tags.append(tag)
        return tags

    # Code used to extract text from headers (h1, h2, etc.)
    headers = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    header_tags = [Tag(name='text') for h in headers]
    for i, h in enumerate(headers):
        header_tags[i].string = h.get_text(strip=True)
        header_tags[i].level = 0

    # Code used to extract text from paragraphs
    paragraphs = soup.find_all('p')
    text_tags = [Tag(name='text') for p in paragraphs]
    for i, p in enumerate(paragraphs):
        text_tags[i].string = p.get_text(strip=True)
        text_tags[i].level = 0

    # Code used to extract text from lists
    lists = soup.find_all(['ul', 'ol'])
    for lst in lists:
        text_tags.extend(process_list_items(lst))

    # Code used to sort the tags based on their order of appearance on the website
    all_tags = header_tags + text_tags
    all_tags.sort(key=lambda x: x.previous_element.index(x) if x.previous_element else 0)

    return '\n'.join(tag.string for tag in all_tags)

def process_text_with_spacy(text):
    # Code used to disable spaCy's named entity recognition (NER) component
    nlp = spacy.load("en_core_web_sm", disable=["ner"])
    doc = nlp(text)
    return doc

def extract_text_spacy(doc):
    return doc.text

def text_to_docx(text, output_path="output.docx"):
    doc = Document()

    # Code used to split the text into lines
    lines = text.split('\n')

    for line in lines:
        # If statement used to skip empty lines
        if line.strip():
            # If statement used to check if the document already has content
            if doc.paragraphs:
                # Code to add a new paragraph for each additional line
                doc.add_paragraph(line)
            else:
                # Code used to add the first line without an additional paragraph
                doc.add_paragraph(line, style='BodyText')

    doc.save(output_path)

if __name__ == "__main__":
    pdf_path = r'C:\\projects\\textdemo.pdf'  # Change this to your PDF file path
    website_url = 'https://none.com'  # Change this to the desired website URL in https://example.com format if you wish to extract text from websites

    # Code used to extract text from PDF
    extracted_text, has_images = pdf_to_text(pdf_path)

    # Code used to extract text from website
    if website_url:
        website_text = extract_text_from_website(website_url)
        extracted_text += "\n" + website_text

    if has_images:
        # Code used to extract text from images
        doc = fitz.open(pdf_path)
        for page_num in range(doc.page_count):
            page = doc[page_num]
            if is_image_page(page):
                image = page.get_pixmap()
                image_text = extract_text_from_image(np.frombuffer(image.samples, dtype=np.uint8).reshape((image.h, image.w, image.n)))
                extracted_text += "\n" + image_text

        doc.close()
        

    # Process the text with spaCy
    spacy_doc = process_text_with_spacy(extracted_text)
    final_extracted_text = extract_text_spacy(spacy_doc)

    # Output to Word document
    output_path = r'C:\\projects\\textdemo.docx'  # Change this to your desired output path
    text_to_docx(final_extracted_text, output_path)

    print(f"Text extracted and saved to {output_path}")

print("Extracted Text:", extracted_text)

